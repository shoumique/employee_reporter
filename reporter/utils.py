"""
Utility functions for the Employee Reporter app.

Handles:
- Bijoy → Unicode conversion (with smart detection to avoid garbling English)
- Excel file loading and processing
- Report preset definitions
- Export Excel generation
- Export DOCX generation (per-employee, zipped when multiple)
"""

import io
import os
import zipfile
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import unicodeconverter as uc
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl.styles import Alignment, Font, PatternFill

# ── Bengali Unicode ranges ─────────────────────────────────────────────────
_BN_START = 0x0980
_BN_END = 0x09FF
# Dependent vowel signs (Mc category)
_VOWEL_SIGN_LO = 0x09BE
_VOWEL_SIGN_HI = 0x09CC
# Khanda Ta: maps to 'r' in Bijoy – invalid mid-word in real Bengali
_KHANDA_TA = "\u09CE"

# ── Fallback column positions used only when auto-detection fails ─────────
_FALLBACK_ID_IDX   = 3   # পার্সোনেল নং  in the standard AGM sheet
_FALLBACK_NAME_IDX = 5   # নাম            in the standard AGM sheet

# Keywords used to locate the ID and name columns by header text
_ID_KEYWORDS   = ["পার্সোনেল", "personnel", "emp_id", "employee_id"]
_ID_EXACT      = ["id"]            # exact (lowercased) column names
_NAME_KEYWORDS = ["নাম", "name"]
_NAME_EXACT    = ["নাম", "name_bn", "name"]


def _detect_id_name_cols(columns: List[str]) -> Tuple[int, int]:
    """Return (id_col_index, name_col_index) by scanning column headers.

    Priority:
    1. Exact match on well-known names (case-insensitive, stripped)
    2. Substring match on keyword lists
    3. Fallback to _FALLBACK_*_IDX (clamped to actual column count)
    """
    lower = [c.lower().strip() for c in columns]

    id_idx: Optional[int] = None
    name_idx: Optional[int] = None

    # ── ID column ──────────────────────────────────────────────────────
    # Exact match first
    for exact in _ID_EXACT:
        if exact in lower:
            id_idx = lower.index(exact)
            break
    # Substring match
    if id_idx is None:
        for i, c in enumerate(lower):
            if any(kw in c for kw in _ID_KEYWORDS):
                id_idx = i
                break
    # Positional fallback
    if id_idx is None:
        id_idx = min(_FALLBACK_ID_IDX, len(columns) - 1)

    # ── Name column ────────────────────────────────────────────────────
    # Exact match first
    for exact in _NAME_EXACT:
        if exact in lower:
            name_idx = lower.index(exact)
            break
    # Substring match
    if name_idx is None:
        for i, c in enumerate(lower):
            if any(kw in c for kw in _NAME_KEYWORDS):
                name_idx = i
                break
    # Positional fallback
    if name_idx is None:
        name_idx = min(_FALLBACK_NAME_IDX, len(columns) - 1)

    # Make sure they're not the same column
    if name_idx == id_idx:
        name_idx = id_idx + 1 if id_idx + 1 < len(columns) else id_idx

    return id_idx, name_idx

# ── Report preset definitions by fixed column positions ───────────────────
REPORT_PRESETS: Dict[str, Dict] = {
    "performance": {
        "label": "Performance Report",
        "description": "Current placement, branch details & performance metrics",
        "icon": "📊",
        "col_positions": [3, 5, 10, 14, 15, 16, 17, 18, 19, 39, 40, 41],
    },
    "appraisal": {
        "label": "Appraisal Report",
        "description": "Educational qualifications & complete promotion history",
        "icon": "🎓",
        "col_positions": [3, 5, 7, 8, 18, 19, 20, 21, 22, 23, 24, 25],
    },
    "basic_info": {
        "label": "Basic Info Report",
        "description": "Personal identification & demographic information",
        "icon": "👤",
        "col_positions": [3, 5, 6, 7, 26, 35, 28, 34],
    },
    "transfer": {
        "label": "Transfer Report",
        "description": "Transfer orders & workplace movement history",
        "icon": "🔄",
        "col_positions": [3, 5, 9, 10, 11, 12, 13, 14, 15, 16, 17],
    },
    "seniority": {
        "label": "Seniority Report",
        "description": "Seniority list with complete promotion timeline",
        "icon": "📋",
        "col_positions": [3, 5, 18, 19, 20, 21, 22, 23, 24, 25, 27, 42],
    },
    "prl": {
        "label": "PRL Report",
        "description": "Pre-retirement leave eligibility & date information",
        "icon": "📅",
        "col_positions": [3, 5, 26, 27, 42, 35],
    },
}


# ── Bijoy → Unicode helpers ────────────────────────────────────────────────

def _has_bengali(text: str) -> bool:
    return any(_BN_START <= ord(c) <= _BN_END for c in text)


def _has_vowel_sign(text: str) -> bool:
    return any(_VOWEL_SIGN_LO <= ord(c) <= _VOWEL_SIGN_HI for c in text)


def _has_invalid_khanda_ta(text: str) -> bool:
    """Khanda Ta mid-word is a strong indicator of garbled English→Bijoy output."""
    for i, ch in enumerate(text):
        if ch == _KHANDA_TA and i + 1 < len(text):
            if _BN_START <= ord(text[i + 1]) <= _BN_END:
                return True
    return False


def convert_bijoy_value(value: Any) -> Any:
    """Convert a cell value from Bijoy to Unicode, with intelligent detection.

    Rules (from main.py in converting_files):
    - Already Unicode Bengali → keep as-is
    - No alphabetic characters (pure numbers/punctuation) → keep as-is
    - After conversion: must have ≥1 vowel sign AND no invalid Khanda Ta → accept
    - Otherwise → return original
    """
    if not isinstance(value, str):
        return value
    text = value.strip()
    if not text:
        return value

    # Pass through already-Unicode Bengali
    if _has_bengali(text):
        return value

    # No alpha → nothing to convert (numbers, symbols)
    if not any(c.isalpha() for c in text):
        return value

    try:
        converted = uc.convert_bijoy_to_unicode(text)
    except Exception:
        return value

    if _has_vowel_sign(converted) and not _has_invalid_khanda_ta(converted):
        return converted

    return value


def _convert_col_name(raw: str) -> str:
    """Convert a single column name from Bijoy to Unicode, cleaning up whitespace.

    Uses the same strict check as convert_bijoy_value (vowel signs required) so
    that plain English identifiers like 'id', 'name_bn', 'designation_en' are
    never accidentally mangled by the Bijoy converter.
    """
    if raw.startswith("Unnamed:"):
        return raw
    clean = raw.replace("\n", " ").strip()
    # Already Unicode Bengali → pass through
    if _has_bengali(clean):
        return clean
    # Pure ASCII with no alpha (numbers, symbols) → nothing to convert
    if not any(c.isalpha() for c in clean):
        return clean
    try:
        converted = uc.convert_bijoy_to_unicode(clean)
        # Require at least one vowel sign, same as cell-value conversion.
        # Real Bijoy-encoded Bengali words always produce vowel signs;
        # English identifiers like 'id' / 'name_bn' produce consonant-only
        # garbage that correctly fails this test.
        if _has_vowel_sign(converted) and not _has_invalid_khanda_ta(converted):
            return converted.replace("\n", " ").strip()
    except Exception:
        pass
    return clean


# ── Core processing ────────────────────────────────────────────────────────

def load_and_process_excel(file_path: str) -> Tuple[pd.DataFrame, List[str]]:
    """
    Load an Excel file, convert Bijoy column names and cell values to Unicode.

    Returns:
        df           – processed DataFrame (unicode column names + values)
        columns      – list of unicode column names (same as df.columns)
    """
    df = pd.read_excel(file_path)

    # Convert column names
    unicode_cols: List[str] = []
    seen: Dict[str, int] = {}
    for col in df.columns:
        name = _convert_col_name(str(col))
        # De-duplicate identical column names
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
        unicode_cols.append(name)

    df.columns = unicode_cols

    # Convert cell values
    df = df.map(lambda v: convert_bijoy_value(v) if isinstance(v, str) else v)

    # Normalise date columns → readable strings so they don't cause serialisation issues
    for col in df.columns:
        if pd.api.types.is_datetime64_any_dtype(df[col]):
            df[col] = df[col].dt.strftime("%d/%m/%Y")

    # Replace NaN with empty string for display
    df = df.where(pd.notna(df), "")

    return df, unicode_cols


def get_employee_list(df: pd.DataFrame) -> List[Dict[str, str]]:
    """Return a list of {id, name} dicts for the employee multi-select."""
    cols = list(df.columns)
    if not cols:
        return []

    id_idx, name_idx = _detect_id_name_cols(cols)
    id_col   = cols[id_idx]
    name_col = cols[name_idx]

    employees = []
    for _, row in df.iterrows():
        emp_id = str(row[id_col]).strip()
        name   = str(row[name_col]).strip()
        if not emp_id or emp_id in ("", "nan"):
            continue
        # Skip placeholder header rows where both ID and name are small integers
        try:
            id_int   = int(emp_id)
            name_int = int(name)
            if 0 < id_int <= 60 and 0 < name_int <= 60:
                continue   # looks like a column-numbering header row
        except (ValueError, TypeError):
            pass
        employees.append({"id": emp_id, "name": name or emp_id})

    return employees


def get_preset_columns(preset_key: str, columns: List[str]) -> List[str]:
    """Return the list of column names for the given preset key."""
    if preset_key not in REPORT_PRESETS:
        return list(columns)
    positions = REPORT_PRESETS[preset_key]["col_positions"]
    return [columns[i] for i in positions if i < len(columns)]


# ── Export ─────────────────────────────────────────────────────────────────

def generate_export_excel(
    df: pd.DataFrame,
    selected_columns: List[str],
    employee_ids: Optional[List[str]] = None,
    report_title: str = "Employee Report",
) -> bytes:
    """
    Build a styled Excel workbook from the processed DataFrame.

    Args:
        df               – processed DataFrame (unicode column names + values)
        selected_columns – columns to include in the output
        employee_ids     – personnel IDs to include; None means all employees
        report_title     – used as the worksheet name and file title

    Returns:
        Raw bytes of the .xlsx file.
    """
    id_idx, _ = _detect_id_name_cols(list(df.columns))
    id_col = list(df.columns)[id_idx]

    # ── Filter rows ──────────────────────────────────────────────────────
    if employee_ids:
        mask = df[id_col].astype(str).isin([str(e).strip() for e in employee_ids])
        df_out = df[mask].copy()
    else:
        df_out = df.copy()

    # ── Filter columns ───────────────────────────────────────────────────
    valid_cols = [c for c in selected_columns if c in df_out.columns]
    if not valid_cols:
        valid_cols = list(df_out.columns)
    df_out = df_out[valid_cols]

    # ── Write to BytesIO ─────────────────────────────────────────────────
    output = io.BytesIO()
    sheet_name = report_title[:31]  # Excel sheet name max 31 chars

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df_out.to_excel(writer, sheet_name=sheet_name, index=False)
        ws = writer.sheets[sheet_name]

        # ── Header styling ────────────────────────────────────────────
        hdr_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
        hdr_font = Font(bold=True, size=11, color="FFFFFF", name="Nirmala UI")
        hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

        for cell in ws[1]:
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = hdr_align

        ws.row_dimensions[1].height = 28

        # ── Body cell font for Bengali ────────────────────────────────
        body_font = Font(name="Nirmala UI", size=10)
        body_align = Alignment(vertical="center", wrap_text=False)

        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.font = body_font
                cell.alignment = body_align

        # ── Auto-size columns ────────────────────────────────────────
        for col_cells in ws.columns:
            max_len = 0
            col_letter = col_cells[0].column_letter
            for cell in col_cells:
                try:
                    cell_len = len(str(cell.value)) if cell.value is not None else 0
                    if cell_len > max_len:
                        max_len = cell_len
                except Exception:
                    pass
            ws.column_dimensions[col_letter].width = max(min(max_len + 4, 42), 12)

        # ── Freeze header row ────────────────────────────────────────
        ws.freeze_panes = "A2"

    output.seek(0)
    return output.read()


# ── DOCX helpers ────────────────────────────────────────────────────────────

def _set_cell_border(cell, **kwargs) -> None:
    """Apply borders to a table cell via direct XML manipulation."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "B0C4DE"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill: str) -> None:
    """Set background fill colour of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)


def _fix_table_col_widths(table, col_twips: tuple) -> None:
    """Enforce fixed column widths on a table via direct XML.

    Sets tblLayout=fixed, patches the existing tblGrid gridCol widths,
    and stamps tcW on every cell so Word cannot widen any column.
    """
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Fixed layout
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # Total table width
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(col_twips)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Patch existing gridCol elements
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for i, gc in enumerate(tblGrid.findall(qn("w:gridCol"))):
            if i < len(col_twips):
                gc.set(qn("w:w"), str(col_twips[i]))

    # Stamp tcW on every cell
    for tbl_row in table.rows:
        for i, cell in enumerate(tbl_row.cells):
            if i < len(col_twips):
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # Remove any existing tcW to avoid duplicates
                for old in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(old)
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(col_twips[i]))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)


def _make_employee_docx(
    row: "pd.Series",
    selected_columns: List[str],
    report_title: str,
) -> bytes:
    """Build a Word document for a single employee.

    Layout:
    - Centred bold title (report_title)
    - 3-column table: Field Name | ঃ | Value
    - Alternating row shading; header-style first column
    """
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Default paragraph style ───────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Nirmala UI"
    style.font.size = Pt(11)

    # ── Header: প্রধান কার্যালয়, ঢাকা ──────────────────────────────────────
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("প্রধান কার্যালয়, ঢাকা")
    header_run.bold           = True
    header_run.font.size      = Pt(13)
    header_run.font.name      = "Nirmala UI"
    header_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ── Title ─────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(report_title)
    title_run.bold      = True
    title_run.font.size = Pt(13)
    title_run.font.name = "Nirmala UI"
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    # Underline the report title to match the original style
    title_run.underline = True

    doc.add_paragraph()  # spacer

    # ── Build rows ────────────────────────────────────────────────────────
    valid_cols = [
        c for c in selected_columns
        if c in row.index and str(row[c]).strip() not in ("", "nan", "NaN")
    ]
    if not valid_cols:
        valid_cols = [c for c in selected_columns if c in row.index]

    # ── Table ─────────────────────────────────────────────────────────────
    table = doc.add_table(rows=len(valid_cols), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # label: 6.5 cm | colon: 0.5 cm | value: 9.0 cm  (twips: 1 cm ≈ 567)
    _COL_TWIPS = (3685, 484, 5102)

    for row_idx, col_name in enumerate(valid_cols):
        cells = table.rows[row_idx].cells

        val = str(row[col_name]).strip()
        if val in ("nan", "NaN"):
            val = ""

        # Cell 0 – field label
        cells[0].text = col_name
        # Cell 1 – separator
        cells[1].text = "ঃ"
        # Cell 2 – value
        cells[2].text = val

        # Row shading: alternate between white and very light blue
        shade = "EEF4FB" if row_idx % 2 == 0 else "FFFFFF"

        for cell_idx, cell in enumerate(cells):
            _set_cell_border(cell)
            _set_cell_shading(cell, shade)

            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = "Nirmala UI"
                    run.font.size = Pt(10)
                    if cell_idx == 0:
                        # Field-name column: slightly bold, dark colour
                        run.bold            = True
                        run.font.color.rgb  = RGBColor(0x1F, 0x4E, 0x79)
                    elif cell_idx == 1:
                        run.bold            = True
                        run.font.color.rgb  = RGBColor(0x44, 0x72, 0xC4)
                    else:
                        run.bold           = False
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Apply fixed column widths after all rows/cells are populated
    _fix_table_col_widths(table, _COL_TWIPS)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def generate_export_docx_zip(
    df: "pd.DataFrame",
    selected_columns: List[str],
    employee_ids: Optional[List[str]] = None,
    report_title: str = "Employee Report",
) -> Tuple[bytes, bool]:
    """Generate one DOCX per employee and optionally wrap them in a ZIP.

    Args:
        df               – processed DataFrame
        selected_columns – columns to show in each document
        employee_ids     – list of personnel IDs to include; None means all
        report_title     – title printed at the top of every document

    Returns:
        (file_bytes, is_zip)
        is_zip=False → file_bytes is a single .docx
        is_zip=True  → file_bytes is a .zip of multiple .docx files
    """
    id_idx, name_idx = _detect_id_name_cols(list(df.columns))
    id_col   = list(df.columns)[id_idx]
    name_col = list(df.columns)[name_idx]

    # ── Filter rows ──────────────────────────────────────────────────────
    if employee_ids:
        mask   = df[id_col].astype(str).isin([str(e).strip() for e in employee_ids])
        df_out = df[mask].copy()
    else:
        df_out = df.copy()

    # ── Filter columns ───────────────────────────────────────────────────
    valid_cols = [c for c in selected_columns if c in df_out.columns]
    if not valid_cols:
        valid_cols = list(df_out.columns)

    # ── Generate per-employee DOCX ───────────────────────────────────────
    docx_files: List[Tuple[str, bytes]] = []
    for _, row in df_out.iterrows():
        emp_id   = str(row[id_col]).strip()
        emp_name = str(row.get(name_col, emp_id)).strip()
        if emp_name in ("nan", "NaN", ""):
            emp_name = emp_id

        docx_bytes = _make_employee_docx(row, valid_cols, report_title)

        safe_name = "".join(
            c if c.isalnum() or c in " _-" else "_" for c in emp_name
        ).strip("_")
        filename = f"{safe_name}_{emp_id}.docx"
        docx_files.append((filename, docx_bytes))

    if not docx_files:
        # Fallback: empty doc
        doc = Document()
        doc.add_paragraph("No data found.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read(), False

    if len(docx_files) == 1:
        return docx_files[0][1], False

    # ── Multiple employees → ZIP ─────────────────────────────────────────
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename, docx_bytes in docx_files:
            zf.writestr(filename, docx_bytes)
    zip_buffer.seek(0)
    return zip_buffer.read(), True
